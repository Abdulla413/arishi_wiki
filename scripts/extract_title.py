import re
from pathlib import Path
from docx import Document

def is_uyghur_title(text):
    text = text.strip()
    if not text: 
        return False
    if '[' in text or ']' in text: 
        return False
    
    words = text.split()
    # Allowing up to 5 words to catch longer numbered titles
    if not (1 <= len(words) <= 5):
        return False

    # UPDATED REGEX:
    # 1. Start with optional digits and punctuation: ^(\d+[\.\،\s\:]*)?
    # 2. Match Uyghur/Arabic script, including Presentation Forms: [\u0600-\u06FF\uFB50-\uFDFF\uFE70-\uFEFF\s]+
    pattern = r'^(\d+[\.\،\s\:]*)?[\u0600-\u06FF\uFB50-\uFDFF\uFE70-\uFEFF\s]+$'
    
    if re.match(pattern, text):
        return True
    return False

def generate_title_list(docx_path, output_txt):
    if not Path(docx_path).exists():
        print(f"❌ Error: File not found at {docx_path}")
        return

    doc = Document(docx_path)
    titles = []
    
    print(f"Checking {len(doc.paragraphs)} paragraphs...")

    for para in doc.paragraphs:
        text = para.text.strip()
        if is_uyghur_title(text):
            titles.append(text)
    
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(titles))
    
    print(f"✅ Extraction complete. {len(titles)} titles saved to {output_txt}")

if __name__ == "__main__":
    SOURCE_DOC = "/media/abdulla-arishi/Volume/arishi.wiki/data/uyghur_tibabiti/uyghurlar_yingi_miwiler.docx"
    LIST_FILE = "titles.txt"
    generate_title_list(SOURCE_DOC, LIST_FILE)
